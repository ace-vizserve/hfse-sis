# Render the academic-head questions document to .docx.
#
# Reads the markdown so there is ONE source of truth — edit the .md, re-run
# this, get a fresh Word file. Handles only the constructs that document
# actually uses: # / ## headings, ---, blockquotes, **bold**, _italic_.
#
# Run: python md_to_docx.py <input.md> <output.docx>
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

INK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x55, 0x55, 0x66)
ACCENT = RGBColor(0x3B, 0x3B, 0x8F)


def shade(el, fill):
    """Paint a paragraph or cell background."""
    pr = el._p.get_or_add_pPr() if hasattr(el, "_p") else el._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:fill"), fill)
    pr.append(s)


def add_runs(par, text, base_bold=False, base_ital=False, color=INK, size=11):
    """Split on **bold** / _italic_ and emit styled runs.

    Recurses into a matched span so nesting works — `**bold with _italic_**`
    would otherwise be emitted as one bold run with the underscores left in."""
    for tok in re.split(r"(\*\*.+?\*\*|_[^_]+?_)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            add_runs(par, tok[2:-2], True, base_ital, color, size)
            continue
        if tok.startswith("_") and tok.endswith("_") and len(tok) > 2:
            add_runs(par, tok[1:-1], base_bold, True, color, size)
            continue
        r = par.add_run(tok)
        r.bold = base_bold
        r.italic = base_ital
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = "Calibri"


def answer_box(doc, label="Your answer:"):
    """A bordered, shaded box she can click into and type."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, "F4F4F8")
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    r.font.name = "Calibri"
    cell.add_paragraph()
    cell.add_paragraph()
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), "C8C8D8")
        borders.append(e)
    tblPr.append(borders)
    doc.add_paragraph()


def main(src, out):
    lines = open(src, encoding="utf-8").read().split("\n")
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)

    quote_buf = []
    para_buf = []

    def flush_para():
        """Markdown hard-wraps; consecutive plain lines are ONE paragraph.
        Joining before rendering also lets **bold** span a line break."""
        if not para_buf:
            return
        body = " ".join(para_buf).strip()
        para_buf.clear()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_runs(p, body)

    def flush_quote():
        """Blockquote → either an answer box or a shaded callout."""
        if not quote_buf:
            return
        body = " ".join(quote_buf).strip()
        quote_buf.clear()
        # Any "Your ...:" prompt is a fill-in box, not a callout — this also
        # catches "Your file / your answer:" on the two send-me-a-sample asks.
        m = re.fullmatch(r"_?(Your [^_]*?:)_?", body)
        if m:
            answer_box(doc, m.group(1))
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        shade(p, "FFF8E8" if "⚠" in body else "F2F4FA")
        add_runs(p, body, size=10.5)

    for raw in lines:
        line = raw.rstrip()

        if line.startswith(">"):
            flush_para()
            quote_buf.append(line.lstrip("> ").rstrip())
            continue
        flush_quote()

        if not line.strip():
            flush_para()
            continue

        if line.startswith(("# ", "## ", "---")):
            flush_para()

        if line.startswith("# "):
            p = doc.add_paragraph()
            add_runs(p, line[2:], base_bold=True, color=ACCENT, size=20)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            add_runs(p, line[3:], base_bold=True, color=ACCENT, size=14)
        elif line.startswith("---"):
            p = doc.add_paragraph()
            pr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:color"), "D8D8E4")
            bd.append(b)
            pr.append(bd)
        else:
            para_buf.append(line.strip())

    flush_para()
    flush_quote()
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
