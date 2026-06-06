"""Generate the staff-facing 'Assumptions to Confirm' Word document with a
fillable answer column. Source of truth for the items is
docs/assumptions-register.md (internal version, keeps the 'if no -> impact'
notes); this produces the shareable .docx Joann / Chandana / officers fill in.

Run:  python scripts/build-assumptions-docx.py
Out:  docs/assumptions-register.docx
"""

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

NAVY = RGBColor(0x21, 0x30, 0x98)
AMBER = RGBColor(0xED, 0x76, 0x22)
HEADER_BG = "213098"
ZEBRA_BG = "EEF1FA"

# (id, who, current behaviour, question, deferred?)
MODULES = [
    (
        "Markbook (grading)",
        "Already confirmed (no need to ask): the grade formula + weights, annual grade, "
        "General Average, the award ladder, and the non-examinable letter grades + per-term "
        "UG/E overrides.",
        [
            ("M1", "Joann",
             "After a grading sheet is locked, changing a grade needs a request where the teacher picks a primary AND secondary approver (office staff); both are notified and a reason is recorded.",
             "When a teacher needs to fix a grade after the sheet is locked, what really happens? Do you pick two approvers, or does one person just approve it directly?",
             False),
            ("M2", "Chandana",
             "For letter-graded subjects (Music, Arts, PE, HE...), the term-to-term trend Alert shows the number-equivalent swing, same as numeric subjects.",
             "For the letter subjects, do you want the term-over-term trend flagged differently, or is showing the number-equivalent fine?",
             True),
            ("M3", "Joann",
             "The publishing checklist is a soft gate: warnings show, but you can always 'Publish anyway'. Nothing blocks publishing.",
             "Before report cards go out, should any of the warnings BLOCK publishing, or is it always your call to publish anyway?",
             False),
        ],
    ),
    (
        "Attendance",
        "Already confirmed: the day-types, vacation leave (1 per term) and compassionate leave "
        "(5 per year) - verified against your actual Term 1 workbook; the HBL / marking-day "
        "overlay matches the published AY2026 calendar.",
        [
            ("A1", "Joann",
             "A student who joins mid-year is only judged from their join date - earlier terms show 0/0 and are never counted against them.",
             "If a student joins in Term 2, should Term 1 attendance just be blank for them - never counted as absences?",
             False),
            ("A2", "Joann",
             "Going over the leave quota shows a warning but still records the leave (you can grant an exception).",
             "If a student goes over their leave quota, should the system let you record it anyway with a warning, or block it?",
             False),
            ("A3", "Teachers",
             "Attendance marks use the system's colours, not the old paper sheet's colours (Present=blue, Absent=yellow, EX=cyan, Late=pink).",
             "Do you want the marks colour-matched to the old paper sheet?",
             True),
        ],
    ),
    (
        "Evaluation (form-class-adviser write-ups)",
        "Already confirmed (the hard way): this module is FCA write-ups only - parent-teacher "
        "conferences run on Ms JoAnn's offline form, not the system. Manual Save-as-draft / Submit.",
        [
            ("E1", "Joann / Chandana",
             "A write-up's text already shows on the report card even as a draft - 'Submit' only drives the progress counts, not whether parents can see it.",
             "Should a comment be hidden from parents until the adviser hits Submit, or is it fine that the draft text already appears?",
             True),
            ("E2", "Chandana",
             "The comment heading reads 'Form Class Adviser's Comments (HFSE Virtues: ...)', pulled from the term's virtue theme.",
             "Is that heading wording right, and is the per-term virtue theme the correct source?",
             False),
            ("E3", "Chandana",
             "The year-end (Term 4) report has no adviser comment block - comments are Term 1-3 only.",
             "Is it correct that the final report has no adviser comment?",
             False),
        ],
    ),
    (
        "Admissions",
        "Already confirmed: the 13-step intake -> 9-stage pipeline; the application-status fields; "
        "STP / Edutrust tracking; and that parents upload ICA documents directly, so the system "
        "doesn't track those files.",
        [
            ("AD1", "Admissions team",
             "After parents upload documents, an officer works through a validation queue approving or rejecting each one.",
             "After a parent uploads documents, who checks them and how? Does one person work a queue approving/rejecting each?",
             False),
            ("AD2", "Admissions team",
             "Chasing splits into 'parent owes us' (To follow / Rejected / Expired) versus 'we owe a review' (Uploaded).",
             "When you chase families for documents, is that the right split - and are those the statuses you actually use?",
             False),
            ("AD3", "Admissions / leadership",
             "There's a page that collects applicant feedback on the online application experience (1-5 rating + comments).",
             "Do you actually ask families to rate the application form? Is this page useful, or unused?",
             False),
            ("AD4", "Joann / Admissions",
             "You can open next year's applications while the current year is still running (the 'early-bird' window), controlled in SIS Admin.",
             "Do you collect next-year applications early while this year runs? Is opening/closing that window from SIS Admin the right place?",
             False),
        ],
    ),
    (
        "Records & SIS Admin (enrolment, identity, settings)",
        "Already confirmed: the permanent student number as the lasting ID; atomic mid-year section "
        "transfer; 'once the school year has started, anyone joining is a late enrollee'; and the "
        "staff-to-role mapping.",
        [
            ("R1", "Joann",
             "Sometimes an enrolled student has no class section yet and won't show in Records until assigned; there's a queue to assign them a section.",
             "Does it happen that a student is enrolled but has no class section (a gap from the old system)? Is assigning them from a queue the right fix?",
             False),
            ("R2", "Joann",
             "The withdrawal reason dropdown (used when a student leaves) has a fixed list of reasons.",
             "Are these the right reasons for why a student leaves - is anything missing?",
             False),
            ("R3", "Joann",
             "A class roll can be re-numbered alphabetically on demand (a button).",
             "Do you ever re-number a class roll alphabetically? Is that something you'd use?",
             False),
            ("R4", "Joann",
             "There's a single feed listing all transfers, withdrawals and late joins together.",
             "Would one combined list of all enrolment changes be useful, or do you track those elsewhere?",
             False),
            ("R5", "Joann / Admissions",
             "Discount codes are managed under SIS Admin (office settings).",
             "Who should own discount codes - the office/SIS Admin, or the admissions team?",
             True),
        ],
    ),
    (
        "P-Files (document renewals)",
        "Already confirmed: it's a document repository (not a review queue); enrolled students only; "
        "and STP documents aren't tracked (parents upload those to ICA directly).",
        [
            ("P1", "P-file officer",
             "When a document is expiring/expired, the officer chases the parent by emailing a reminder or marking 'promised by <date>', with a 24-hour gap between reminders.",
             "When a passport/pass is expiring, how do you chase the family? Is emailing a reminder / marking a promised date how you'd work?",
             False),
            ("P2", "P-file officer",
             "'Expiring soon' surfaces at 30 / 60 / 90 days before the expiry date.",
             "How far ahead do you start chasing renewals - 30, 60, 90 days?",
             False),
        ],
    ),
    (
        "Parent portal",
        "Already confirmed: parents see report cards gated by publication windows; secure sign-in.",
        [
            ("PA1", "Joann",
             "Saving a report card as a PDF is done through the browser's Print (there's no one-click PDF download button).",
             "Is 'Print -> Save as PDF' acceptable for report cards, or do you need a one-click PDF download?",
             True),
            ("PA2", "Joann",
             "Report cards are released per class-and-term, on a publication window you open and close.",
             "Is releasing report cards per class-and-term, on a window you open, the right model?",
             False),
        ],
    ),
]


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9, italic=False, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color


def set_col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C7CEE6")
        borders.append(el)
    tblPr.append(borders)


doc = Document()

# Landscape, narrow margins for table room.
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)

# Title
t = doc.add_paragraph()
tr = t.add_run("HFSE SIS - Assumptions to Confirm")
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sr = sub.add_run(
    "The system's academic core (grading, attendance quotas, the admissions pipeline) is "
    "already confirmed. This checklist covers the workflow assumptions we'd like you to check. "
    "For each row: if what the system does matches how you actually work, write \"Correct\" in "
    "the last column. If it's different, please describe what you really do. Each question notes "
    "who it's mainly for - feel free to add comments on any row."
)
sr.font.size = Pt(10)

legend = doc.add_paragraph()
lr = legend.add_run("Rows marked ⏸ are decisions already waiting on your input. "
                    "The highest-value ones to get right are M1, AD1 and P1 - they shape how staff spend their day.")
lr.font.size = Pt(9)
lr.italic = True
lr.font.color.rgb = AMBER

WIDTHS = [0.35, 0.95, 2.75, 2.95, 2.55]  # ~9.55" usable in landscape

for module_name, confirmed, items in MODULES:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    hr = h.add_run(module_name)
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = NAVY

    c = doc.add_paragraph()
    cr = c.add_run(confirmed)
    cr.italic = True
    cr.font.size = Pt(8.5)
    cr.font.color.rgb = RGBColor(0x55, 0x5B, 0x6E)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_borders(table)
    hdr = table.rows[0].cells
    for i, label in enumerate(["#", "For", "What the system does today", "Question for you", "Your answer"]):
        set_cell_text(hdr[i], label, bold=True, size=9, white=True)
        shade(hdr[i], HEADER_BG)

    for idx, (iid, who, current, question, deferred) in enumerate(items):
        cells = table.add_row().cells
        set_cell_text(cells[0], iid + (" ⏸" if deferred else ""), bold=True, size=8.5)
        set_cell_text(cells[1], who, size=8.5)
        set_cell_text(cells[2], current, size=8.5)
        set_cell_text(cells[3], question, size=8.5, color=NAVY)
        set_cell_text(cells[4], "", size=9)  # empty fill-in box
        if idx % 2 == 1:
            for ci in (0, 1, 2, 3):
                shade(cells[ci], ZEBRA_BG)
    set_col_widths(table, WIDTHS)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    "Thank you. Anything you mark as different becomes a small change to the system - "
    "most are simplifications. Hand this back once filled in (or add comments inline)."
)
fr.font.size = Pt(9)
fr.italic = True

import os
out = os.path.join("docs", "assumptions-register.docx")
doc.save(out)
print(f"wrote {out}")
